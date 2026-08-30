"""
================================================================================
STAGE 6: DOCX AND PDF OUTPUT
================================================================================

Classes:
- GenerateDocument: Render the tailored CV (Stage 5 output) plus the original
  contact/education details (Stage 1 output) into a .docx and a .pdf file.

Imported by main.py

IMPORT:
from step_6_docx import GenerateDocument

================================================================================
"""

import re

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet


def _safe_filename(name, company):
    """Build a filesystem-safe filename from candidate name and company"""
    base = f"CV_{name}_{company}".strip("_")
    return re.sub(r"[^A-Za-z0-9_-]+", "_", base) or "CV"


class GenerateDocument:
    """Render tailored CV content to .docx and .pdf"""

    def __init__(self):
        """Initialize"""
        print("[*] Initializing document generator (python-docx + reportlab)...")
        print("[✓] Generator ready")

    def run(self, resume_data, job_data, generated_cv, output_dir="output"):
        """Generate .docx and .pdf, return their paths"""
        print("\n" + "=" * 70)
        print("STAGE 6: GENERATE DOCX + PDF")
        print("=" * 70)

        filename = _safe_filename(
            resume_data.get("name", "Candidate"), job_data.get("company", "")
        )
        docx_path = f"{output_dir}/{filename}.docx"
        pdf_path = f"{output_dir}/{filename}.pdf"

        self._build_docx(resume_data, job_data, generated_cv, docx_path)
        print(f"[✓] Saved: {docx_path}")

        self._build_pdf(resume_data, job_data, generated_cv, pdf_path)
        print(f"[✓] Saved: {pdf_path}")

        return {"docx_path": docx_path, "pdf_path": pdf_path}

    @staticmethod
    def _build_docx(resume_data, job_data, generated_cv, path):
        doc = Document()

        name_heading = doc.add_heading(resume_data.get("name", ""), level=0)
        name_heading.alignment = 1

        contact_bits = [
            resume_data.get("email", ""),
            resume_data.get("phone", ""),
            resume_data.get("location", ""),
        ]
        contact_line = doc.add_paragraph(" | ".join(b for b in contact_bits if b))
        contact_line.alignment = 1

        doc.add_heading("Professional Summary", level=1)
        doc.add_paragraph(generated_cv.get("summary", ""))

        doc.add_heading(
            f"Experience — tailored for {job_data.get('title', '')} "
            f"at {job_data.get('company', '')}",
            level=1,
        )
        experience = (resume_data.get("experience") or [{}])[0]
        doc.add_paragraph(
            f"{experience.get('position', '')} — {experience.get('company', '')} "
            f"({experience.get('duration', '')})"
        ).runs[0].bold = True
        for bullet in generated_cv.get("bullets", []):
            doc.add_paragraph(bullet, style="List Bullet")

        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(generated_cv.get("ordered_skills", [])))

        if resume_data.get("education"):
            doc.add_heading("Education", level=1)
            for edu in resume_data["education"]:
                line = " ".join(
                    filter(None, [edu.get("degree", ""), edu.get("field", "")])
                )
                doc.add_paragraph(
                    f"{line} — {edu.get('school', '')} ({edu.get('graduation_year', '')})"
                )

        for style_name in ("Normal",):
            doc.styles[style_name].font.size = Pt(11)

        doc.save(path)

    @staticmethod
    def _build_pdf(resume_data, job_data, generated_cv, path):
        styles = getSampleStyleSheet()
        doc = SimpleDocTemplate(
            path, pagesize=LETTER,
            leftMargin=0.75 * inch, rightMargin=0.75 * inch,
            topMargin=0.75 * inch, bottomMargin=0.75 * inch,
        )

        story = [
            Paragraph(resume_data.get("name", ""), styles["Title"]),
            Paragraph(
                " | ".join(filter(None, [
                    resume_data.get("email", ""),
                    resume_data.get("phone", ""),
                    resume_data.get("location", ""),
                ])),
                styles["Normal"],
            ),
            Spacer(1, 0.2 * inch),
            Paragraph("Professional Summary", styles["Heading2"]),
            Paragraph(generated_cv.get("summary", ""), styles["Normal"]),
            Spacer(1, 0.15 * inch),
            Paragraph(
                f"Experience — tailored for {job_data.get('title', '')} "
                f"at {job_data.get('company', '')}",
                styles["Heading2"],
            ),
        ]

        experience = (resume_data.get("experience") or [{}])[0]
        story.append(Paragraph(
            f"<b>{experience.get('position', '')} — {experience.get('company', '')} "
            f"({experience.get('duration', '')})</b>",
            styles["Normal"],
        ))
        story.append(ListFlowable(
            [ListItem(Paragraph(b, styles["Normal"])) for b in generated_cv.get("bullets", [])],
            bulletType="bullet",
        ))

        story.append(Spacer(1, 0.15 * inch))
        story.append(Paragraph("Skills", styles["Heading2"]))
        story.append(Paragraph(", ".join(generated_cv.get("ordered_skills", [])), styles["Normal"]))

        if resume_data.get("education"):
            story.append(Spacer(1, 0.15 * inch))
            story.append(Paragraph("Education", styles["Heading2"]))
            for edu in resume_data["education"]:
                line = " ".join(filter(None, [edu.get("degree", ""), edu.get("field", "")]))
                story.append(Paragraph(
                    f"{line} — {edu.get('school', '')} ({edu.get('graduation_year', '')})",
                    styles["Normal"],
                ))

        doc.build(story)
