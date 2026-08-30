"""Render a generated CV (no target job, unlike render_tailored_document.py)
to .docx and .pdf, in the shared ATS-friendly format."""

from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem

import document_style as style


class ProfessionalDocumentGenerator:
    """Render generated CV content to .docx and .pdf"""

    def run(self, profile_data, generated_cv, output_dir="output/core"):
        filename = style.safe_filename(profile_data.get("name", "Candidate"))
        docx_path = f"{output_dir}/{filename}.docx"
        pdf_path = f"{output_dir}/{filename}.pdf"

        self._build_docx(profile_data, generated_cv, docx_path)
        self._build_pdf(profile_data, generated_cv, pdf_path)

        return {"docx_path": docx_path, "pdf_path": pdf_path}

    @staticmethod
    def _build_docx(profile_data, generated_cv, path):
        doc = Document()
        style.set_base_style(doc)
        style.add_name_heading(doc, profile_data.get("name", ""))

        contact = [profile_data.get("email", ""), profile_data.get("phone", ""), profile_data.get("location", "")]
        doc.add_paragraph(" | ".join(b for b in contact if b))

        style.add_section_heading(doc, "Professional Summary")
        doc.add_paragraph(generated_cv.get("summary", ""))

        style.add_section_heading(doc, "Skills")
        doc.add_paragraph(", ".join(generated_cv.get("ordered_skills", [])))

        style.add_section_heading(doc, "Experience")
        for job in generated_cv.get("experience", []):
            title_para = doc.add_paragraph()
            title_para.add_run(f"{job.get('position', '')} - {job.get('company', '')}").bold = True
            title_para.add_run(f"   ({style.date_range(job.get('duration', ''))})").italic = True
            for bullet in job.get("bullets", []):
                doc.add_paragraph(bullet, style="List Bullet")

        if profile_data.get("education"):
            style.add_section_heading(doc, "Education")
            for edu in profile_data["education"]:
                line = " ".join(filter(None, [edu.get("degree", ""), edu.get("field", "")]))
                doc.add_paragraph(f"{line} - {edu.get('school', '')} ({edu.get('graduation_year', '')})")

        doc.save(path)

    @staticmethod
    def _build_pdf(profile_data, generated_cv, path):
        styles = style.pdf_styles()
        doc = SimpleDocTemplate(
            path, pagesize=LETTER,
            leftMargin=0.75 * inch, rightMargin=0.75 * inch,
            topMargin=0.75 * inch, bottomMargin=0.75 * inch,
        )

        contact = [profile_data.get("email", ""), profile_data.get("phone", ""), profile_data.get("location", "")]
        story = [
            Paragraph(profile_data.get("name", ""), styles["ATSName"]),
            Paragraph(" | ".join(filter(None, contact)), styles["Normal"]),
            Paragraph("PROFESSIONAL SUMMARY", styles["ATSHeading"]),
            Paragraph(generated_cv.get("summary", ""), styles["Normal"]),
            Paragraph("SKILLS", styles["ATSHeading"]),
            Paragraph(", ".join(generated_cv.get("ordered_skills", [])), styles["Normal"]),
            Paragraph("EXPERIENCE", styles["ATSHeading"]),
        ]

        for job in generated_cv.get("experience", []):
            story.append(Paragraph(
                f"<b>{job.get('position', '')} - {job.get('company', '')}</b>"
                f"  <i>({style.date_range(job.get('duration', ''))})</i>",
                styles["Normal"],
            ))
            if job.get("bullets"):
                story.append(ListFlowable(
                    [ListItem(Paragraph(b, styles["Normal"])) for b in job["bullets"]],
                    bulletType="bullet",
                ))
            story.append(Spacer(1, 0.08 * inch))

        if profile_data.get("education"):
            story.append(Paragraph("EDUCATION", styles["ATSHeading"]))
            for edu in profile_data["education"]:
                line = " ".join(filter(None, [edu.get("degree", ""), edu.get("field", "")]))
                story.append(Paragraph(
                    f"{line} - {edu.get('school', '')} ({edu.get('graduation_year', '')})",
                    styles["Normal"],
                ))

        doc.build(story)
