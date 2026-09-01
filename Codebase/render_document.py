"""Render generated CV content (summary/experience/ordered_skills) to a
.docx and a .pdf, in the shared ATS-friendly format. Used by both the
tailored-CV pipeline and the profile-CV pipeline - they produce the same
CV data shape, differing only in the experience heading and filename.
"""

from docx import Document
from reportlab.lib.colors import black
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem, HRFlowable

import document_style as style


class DocumentGenerator:
    """Render generated CV content to .docx and .pdf"""

    def run(self, person_data, generated_cv, output_dir, experience_heading="Experience", filename_suffix=""):
        filename = style.safe_filename(person_data.get("name", "Candidate"), filename_suffix)
        docx_path = f"{output_dir}/{filename}.docx"
        pdf_path = f"{output_dir}/{filename}.pdf"

        self._build_docx(person_data, generated_cv, docx_path, experience_heading)
        self._build_pdf(person_data, generated_cv, pdf_path, experience_heading)

        return {"docx_path": docx_path, "pdf_path": pdf_path}

    @staticmethod
    def _build_docx(person_data, generated_cv, path, experience_heading):
        doc = Document()
        style.set_base_style(doc)
        style.add_name_heading(doc, person_data.get("name", ""))

        contact = [person_data.get("email", ""), person_data.get("phone", ""), person_data.get("location", "")]
        doc.add_paragraph(" | ".join(b for b in contact if b))

        style.add_section_heading(doc, "Professional Summary")
        doc.add_paragraph(generated_cv.get("summary", ""))

        style.add_section_heading(doc, "Skills")
        doc.add_paragraph(", ".join(generated_cv.get("ordered_skills", [])))

        style.add_section_heading(doc, experience_heading)
        for job in generated_cv.get("experience", []):
            title_para = doc.add_paragraph()
            title_para.add_run(f"{job.get('position', '')} - {job.get('company', '')}").bold = True
            title_para.add_run(f"   ({style.date_range(job.get('duration', ''))})").italic = True
            for bullet in job.get("bullets", []):
                doc.add_paragraph(bullet, style="List Bullet")

        if person_data.get("education"):
            style.add_section_heading(doc, "Education")
            for edu in person_data["education"]:
                line = " ".join(filter(None, [edu.get("degree", ""), edu.get("field", "")]))
                doc.add_paragraph(f"{line} - {edu.get('school', '')} ({edu.get('graduation_year', '')})")

        if person_data.get("certifications"):
            style.add_section_heading(doc, "Certifications")
            for cert in person_data["certifications"]:
                doc.add_paragraph(style.format_certification(cert), style="List Bullet")

        if person_data.get("languages"):
            style.add_section_heading(doc, "Languages")
            doc.add_paragraph(", ".join(person_data["languages"]))

        doc.save(path)

    @staticmethod
    def _build_pdf(person_data, generated_cv, path, experience_heading):
        styles = style.pdf_styles()
        doc = SimpleDocTemplate(
            path, pagesize=LETTER,
            leftMargin=0.75 * inch, rightMargin=0.75 * inch,
            topMargin=0.75 * inch, bottomMargin=0.75 * inch,
        )

        def heading(text):
            """Bold heading plus a bottom rule, mirroring the Word renderer's
            bordered section headings so the two outputs read the same way."""
            return [
                Paragraph(text.upper(), styles["ATSHeading"]),
                HRFlowable(width="100%", thickness=0.75, color=black, spaceAfter=6),
            ]

        contact = [person_data.get("email", ""), person_data.get("phone", ""), person_data.get("location", "")]
        story = [
            Paragraph(person_data.get("name", ""), styles["ATSName"]),
            Paragraph(" | ".join(filter(None, contact)), styles["ATSContact"]),
            *heading("Professional Summary"),
            Paragraph(generated_cv.get("summary", ""), styles["Normal"]),
            *heading("Skills"),
            Paragraph(", ".join(generated_cv.get("ordered_skills", [])), styles["Normal"]),
            *heading(experience_heading),
        ]

        for job in generated_cv.get("experience", []):
            story.append(Paragraph(
                f"<b>{job.get('position', '')} - {job.get('company', '')}</b>"
                f"  <i>({style.date_range(job.get('duration', ''))})</i>",
                styles["Normal"],
            ))
            if job.get("bullets"):
                story.append(ListFlowable(
                    [ListItem(Paragraph(b, styles["Normal"])) for b in job["bullets"]],
                    bulletType="bullet",
                ))
            story.append(Spacer(1, 0.1 * inch))

        if person_data.get("education"):
            story.extend(heading("Education"))
            for edu in person_data["education"]:
                line = " ".join(filter(None, [edu.get("degree", ""), edu.get("field", "")]))
                story.append(Paragraph(
                    f"{line} - {edu.get('school', '')} ({edu.get('graduation_year', '')})",
                    styles["Normal"],
                ))

        if person_data.get("certifications"):
            story.extend(heading("Certifications"))
            story.append(ListFlowable(
                [ListItem(Paragraph(style.format_certification(c), styles["Normal"])) for c in person_data["certifications"]],
                bulletType="bullet",
            ))

        if person_data.get("languages"):
            story.extend(heading("Languages"))
            story.append(Paragraph(", ".join(person_data["languages"]), styles["Normal"]))

        doc.build(story)
