"""
================================================================================
CORE TASK - STEP 4: DOCX AND PDF OUTPUT (ATS-FRIENDLY FORMAT)
================================================================================

Classes:
- ProfessionalDocumentGenerator: Render a generated CV (no target job, unlike
  step_6_docx.py's tailored output) to a .docx and a .pdf file, following an
  ATS-friendly layout:
    - single column, no tables/text-boxes/columns/graphics
    - one standard sans-serif font throughout (Calibri / Helvetica)
    - plain black text, plain-text section headings (no theme colors)
    - real bulleted lists, not unicode bullet characters typed into text
    - reverse-chronological experience, consistent "Start - End" dates
    - text-based PDF (not a scanned image), so both files stay parseable

Imported by core_main.py

IMPORT:
from core_4_docx import ProfessionalDocumentGenerator

================================================================================
"""

import re

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.lib.colors import black
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem

BODY_FONT = "Calibri"
BLACK = RGBColor(0x00, 0x00, 0x00)


def _safe_filename(name):
    base = f"CV_{name}".strip("_")
    return re.sub(r"[^A-Za-z0-9_-]+", "_", base) or "CV"


def _range(duration):
    """Normalise a duration string to 'Start - End' (plain hyphen, ATS-safe)"""
    return re.sub(r"\s*(?:-|to|–|—)\s*", " - ", duration or "").strip()


class ProfessionalDocumentGenerator:
    """Render generated CV content to an ATS-friendly .docx and .pdf"""

    def run(self, profile_data, generated_cv, output_dir="output/core"):
        filename = _safe_filename(profile_data.get("name", "Candidate"))
        docx_path = f"{output_dir}/{filename}.docx"
        pdf_path = f"{output_dir}/{filename}.pdf"

        self._build_docx(profile_data, generated_cv, docx_path)
        self._build_pdf(profile_data, generated_cv, pdf_path)

        return {"docx_path": docx_path, "pdf_path": pdf_path}

    # ------------------------------------------------------------------
    # DOCX
    # ------------------------------------------------------------------

    @staticmethod
    def _set_base_style(doc):
        normal = doc.styles["Normal"]
        normal.font.name = BODY_FONT
        normal.font.size = Pt(11)
        normal.font.color.rgb = BLACK
        # Word only honours the ascii font id unless the east-asian id is
        # also set, so pin both to avoid a silent font substitution.
        normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)

    @staticmethod
    def _add_section_heading(doc, text):
        """Plain bold black heading with a bottom rule - no theme colors,
        so it stays legible after an ATS strips all formatting to text."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = BLACK
        run.font.name = BODY_FONT

        p_border = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        p_border.append(bottom)
        p.paragraph_format.element.get_or_add_pPr().append(p_border)
        return p

    def _build_docx(self, profile_data, generated_cv, path):
        doc = Document()
        self._set_base_style(doc)

        for section in doc.sections:
            section.left_margin = Pt(54)
            section.right_margin = Pt(54)

        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name_run = name_para.add_run(profile_data.get("name", ""))
        name_run.bold = True
        name_run.font.size = Pt(18)
        name_run.font.color.rgb = BLACK
        name_run.font.name = BODY_FONT

        contact_bits = [
            profile_data.get("email", ""),
            profile_data.get("phone", ""),
            profile_data.get("location", ""),
        ]
        doc.add_paragraph(" | ".join(b for b in contact_bits if b))

        self._add_section_heading(doc, "Professional Summary")
        doc.add_paragraph(generated_cv.get("summary", ""))

        self._add_section_heading(doc, "Skills")
        doc.add_paragraph(", ".join(generated_cv.get("ordered_skills", [])))

        self._add_section_heading(doc, "Experience")
        for job in generated_cv.get("experience", []):
            title_para = doc.add_paragraph()
            title_para.paragraph_format.space_after = Pt(2)
            run = title_para.add_run(f"{job.get('position', '')} - {job.get('company', '')}")
            run.bold = True
            date_run = title_para.add_run(f"   ({_range(job.get('duration', ''))})")
            date_run.italic = True
            for bullet in job.get("bullets", []):
                bullet_para = doc.add_paragraph(bullet, style="List Bullet")
                bullet_para.paragraph_format.space_after = Pt(2)

        if profile_data.get("education"):
            self._add_section_heading(doc, "Education")
            for edu in profile_data["education"]:
                line = " ".join(filter(None, [edu.get("degree", ""), edu.get("field", "")]))
                doc.add_paragraph(
                    f"{line} - {edu.get('school', '')} ({edu.get('graduation_year', '')})"
                )

        doc.save(path)

    # ------------------------------------------------------------------
    # PDF
    # ------------------------------------------------------------------

    @staticmethod
    def _pdf_styles():
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            "ATSName", parent=styles["Normal"],
            fontName="Helvetica-Bold", fontSize=18, textColor=black, spaceAfter=4,
        ))
        styles.add(ParagraphStyle(
            "ATSHeading", parent=styles["Normal"],
            fontName="Helvetica-Bold", fontSize=12, textColor=black,
            spaceBefore=14, spaceAfter=4, borderPadding=0,
        ))
        styles["Normal"].fontName = "Helvetica"
        styles["Normal"].fontSize = 11
        styles["Normal"].textColor = black
        return styles

    def _build_pdf(self, profile_data, generated_cv, path):
        styles = self._pdf_styles()
        doc = SimpleDocTemplate(
            path, pagesize=LETTER,
            leftMargin=0.75 * inch, rightMargin=0.75 * inch,
            topMargin=0.75 * inch, bottomMargin=0.75 * inch,
        )

        story = [
            Paragraph(profile_data.get("name", ""), styles["ATSName"]),
            Paragraph(
                " | ".join(filter(None, [
                    profile_data.get("email", ""),
                    profile_data.get("phone", ""),
                    profile_data.get("location", ""),
                ])),
                styles["Normal"],
            ),
            Paragraph("PROFESSIONAL SUMMARY", styles["ATSHeading"]),
            Paragraph(generated_cv.get("summary", ""), styles["Normal"]),
            Paragraph("SKILLS", styles["ATSHeading"]),
            Paragraph(", ".join(generated_cv.get("ordered_skills", [])), styles["Normal"]),
            Paragraph("EXPERIENCE", styles["ATSHeading"]),
        ]

        for job in generated_cv.get("experience", []):
            story.append(Paragraph(
                f"<b>{job.get('position', '')} - {job.get('company', '')}</b>"
                f"  <i>({_range(job.get('duration', ''))})</i>",
                styles["Normal"],
            ))
            if job.get("bullets"):
                story.append(ListFlowable(
                    [ListItem(Paragraph(b, styles["Normal"])) for b in job["bullets"]],
                    bulletType="bullet",
                ))
            story.append(Spacer(1, 0.08 * inch))

        if profile_data.get("education"):
            story.append(Paragraph("EDUCATION", styles["ATSHeading"]))
            for edu in profile_data["education"]:
                line = " ".join(filter(None, [edu.get("degree", ""), edu.get("field", "")]))
                story.append(Paragraph(
                    f"{line} - {edu.get('school', '')} ({edu.get('graduation_year', '')})",
                    styles["Normal"],
                ))

        doc.build(story)
