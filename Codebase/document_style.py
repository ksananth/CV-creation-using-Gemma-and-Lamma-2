"""Shared ATS-friendly styling for the .docx and .pdf CV renderers.

Single column, one standard sans-serif font, plain black text, plain-text
section headings (no theme colors), real bulleted lists - so both the Word
and PDF output stay parseable by applicant tracking systems.
"""

import re

from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.colors import black
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

BODY_FONT = "Calibri"
BLACK = RGBColor(0x00, 0x00, 0x00)


def safe_filename(*parts):
    base = "CV_" + "_".join(p for p in parts if p)
    return re.sub(r"[^A-Za-z0-9_-]+", "_", base).strip("_") or "CV"


def date_range(duration):
    """Normalise a duration string to 'Start - End' (plain hyphen, ATS-safe)"""
    return re.sub(r"\s*(?:-|to|–|—)\s*", " - ", duration or "").strip()


def set_base_style(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    for section in doc.sections:
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)


def add_name_heading(doc, name):
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = BLACK
    run.font.name = BODY_FONT
    return p


def add_section_heading(doc, text):
    """Bold black heading with a bottom rule - no theme colors, so it stays
    legible after an ATS strips formatting down to plain text."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = BLACK
    run.font.name = BODY_FONT

    border = OxmlElement("w:bottom")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), "6")
    border.set(qn("w:space"), "1")
    border.set(qn("w:color"), "000000")
    p_border = OxmlElement("w:pBdr")
    p_border.append(border)
    p.paragraph_format.element.get_or_add_pPr().append(p_border)
    return p


def format_certification(item):
    """A certification entry may be a plain string or a structured dict
    (certification/issued/valid_through) - local models don't always follow
    the simpler schema, so render either shape rather than dropping it."""
    if isinstance(item, dict):
        name = item.get("certification") or item.get("name") or ""
        issued = item.get("issued", "")
        valid = item.get("valid_through", "")
        dates = f"{issued} - {valid}" if issued and valid else (issued or valid)
        return f"{name} ({dates})" if dates else name
    return str(item)


def pdf_styles():
    """Build the paragraph styles used by the PDF renderer.

    Every style sets its own `leading` explicitly instead of relying on
    inheritance from Normal: ParagraphStyle copies parent attributes at
    construction time, not dynamically, so a style built before Normal's
    fontSize/leading were customized would keep Normal's original 10pt
    leading - too short for an 18pt heading, causing the text to visually
    overlap the line below it.
    """
    styles = getSampleStyleSheet()
    styles["Normal"].fontName = "Helvetica"
    styles["Normal"].fontSize = 11
    styles["Normal"].leading = 15
    styles["Normal"].textColor = black
    styles["Normal"].spaceAfter = 6

    styles.add(ParagraphStyle(
        "ATSName", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=18, leading=22, textColor=black,
        spaceAfter=2,
    ))
    styles.add(ParagraphStyle(
        "ATSContact", parent=styles["Normal"],
        spaceAfter=12,
    ))
    styles.add(ParagraphStyle(
        "ATSHeading", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=12, leading=15, textColor=black,
        spaceBefore=14, spaceAfter=2,
    ))
    return styles
